"""
Document text extraction for lecture materials (PDF, DOCX)
"""
import os
from pathlib import Path
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document


class PDFReader:
    """Handles PDF and DOCX text extraction"""
    
    def __init__(self):
        """Initialize the document reader"""
        pass
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract all text content from a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text content from all pages
        """
        if not os.path.isfile(pdf_path):
            raise ValueError(f"Invalid PDF file path: {pdf_path}")
        
        if not pdf_path.lower().endswith('.pdf'):
            raise ValueError(f"File is not a PDF: {pdf_path}")
        
        try:
            # Open and read the PDF
            reader = PdfReader(pdf_path)
            
            if len(reader.pages) == 0:
                raise ValueError("PDF has no pages")
            
            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    # Add page marker for reference
                    text_content.append(f"=== Page {page_num} ===\n{page_text.strip()}")
            
            if not text_content:
                raise ValueError("No text could be extracted from PDF")
            
            # Join all pages with double newline
            full_text = "\n\n".join(text_content)
            
            return full_text
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")
    
    def get_pdf_info(self, pdf_path: str) -> dict:
        """
        Get metadata information about a PDF
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary with PDF metadata
        """
        try:
            reader = PdfReader(pdf_path)
            metadata = reader.metadata or {}
            
            return {
                "num_pages": len(reader.pages),
                "title": metadata.get("/Title", "Unknown"),
                "author": metadata.get("/Author", "Unknown"),
                "subject": metadata.get("/Subject", "Unknown"),
                "creator": metadata.get("/Creator", "Unknown"),
            }
        except Exception as e:
            return {
                "error": f"Failed to read PDF metadata: {str(e)}"
            }
    
    def validate_pdf(self, pdf_path: str) -> bool:
        """
        Check if a file is a valid PDF
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            True if valid PDF, False otherwise
        """
        try:
            if not os.path.isfile(pdf_path):
                return False
            
            if not pdf_path.lower().endswith('.pdf'):
                return False
            
            # Try to open it
            reader = PdfReader(pdf_path)
            
            # Check if it has at least one page
            if len(reader.pages) == 0:
                return False
            
            return True
        except Exception:
            return False
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract all text content from a DOCX file
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content from all paragraphs and tables
        """
        if not os.path.isfile(docx_path):
            raise ValueError(f"Invalid DOCX file path: {docx_path}")
        
        if not docx_path.lower().endswith('.docx'):
            raise ValueError(f"File is not a DOCX: {docx_path}")
        
        try:
            # Open and read the DOCX
            doc = Document(docx_path)
            
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                raise ValueError("DOCX has no content")
            
            # Extract text from all paragraphs
            text_content = []
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                if para_text:
                    text_content.append(para_text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise ValueError("No text could be extracted from DOCX")
            
            # Join all content with double newline
            full_text = "\n\n".join(text_content)
            
            return full_text
            
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")
    
    def validate_docx(self, docx_path: str) -> bool:
        """
        Check if a file is a valid DOCX
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            True if valid DOCX, False otherwise
        """
        try:
            if not os.path.isfile(docx_path):
                return False
            
            if not docx_path.lower().endswith('.docx'):
                return False
            
            # Try to open it
            doc = Document(docx_path)
            
            # Check if it has any content
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                return False
            
            return True
        except Exception:
            return False

